#!/usr/bin/env python3
"""Genera el manual de usuario en formato Word (.docx)."""

from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm, Pt, RGBColor

RAIZ = Path(__file__).resolve().parents[1]
SALIDA = RAIZ / "Manual-de-usuario-ERP-Tienda.docx"

COLOR_PRIMARIO = RGBColor(0x1A, 0x3A, 0x6B)
COLOR_ACENTO = RGBColor(0x2B, 0x6C, 0xB0)
COLOR_TEXTO = RGBColor(0x2C, 0x3E, 0x50)


def configurar_estilos(documento: Document) -> None:
    estilos = documento.styles
    normal = estilos["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = COLOR_TEXTO
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    normal.paragraph_format.line_spacing = 1.15
    normal.paragraph_format.space_after = Pt(6)

    for nivel, tamano in [(1, 18), (2, 14), (3, 12)]:
        estilo = estilos[f"Heading {nivel}"]
        estilo.font.name = "Calibri Light"
        estilo.font.bold = True
        estilo.font.size = Pt(tamano)
        estilo.font.color.rgb = COLOR_PRIMARIO if nivel == 1 else COLOR_ACENTO
        estilo.paragraph_format.space_before = Pt(14 if nivel == 1 else 10)
        estilo.paragraph_format.space_after = Pt(8)


def agregar_pie_pagina(documento: Document, texto: str) -> None:
    seccion = documento.sections[0]
    pie = seccion.footer
    parrafo = pie.paragraphs[0] if pie.paragraphs else pie.add_paragraph()
    parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    parrafo.text = texto
    for run in parrafo.runs:
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x7F, 0x8C, 0x8D)


def agregar_campo_numero_pagina(parrafo) -> None:
    run = parrafo.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_char_sep = OxmlElement("w:fldChar")
    fld_char_sep.set(qn("w:fldCharType"), "separate")
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_begin)
    run._r.append(instr)
    run._r.append(fld_char_sep)
    run._r.append(fld_char_end)


def portada(documento: Document) -> None:
    for _ in range(6):
        documento.add_paragraph()

    titulo = documento.add_paragraph()
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = titulo.add_run("ERP Tienda")
    run.bold = True
    run.font.size = Pt(32)
    run.font.color.rgb = COLOR_PRIMARIO
    run.font.name = "Calibri Light"

    subt = documento.add_paragraph()
    subt.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = subt.add_run("Manual de usuario")
    run2.font.size = Pt(20)
    run2.font.color.rgb = COLOR_ACENTO

    documento.add_paragraph()
    desc = documento.add_paragraph()
    desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = desc.add_run(
        "Guía operativa para el uso diario del sistema:\n"
        "clientes, proveedores, productos, ventas, compras, stock y reportes."
    )
    r3.font.size = Pt(12)
    r3.font.color.rgb = COLOR_TEXTO

    documento.add_paragraph()
    fecha = documento.add_paragraph()
    fecha.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rf = fecha.add_run(f"Versión {date.today().strftime('%B %Y')}")
    rf.italic = True
    rf.font.size = Pt(11)
    rf.font.color.rgb = RGBColor(0x95, 0xA5, 0xA6)

    documento.add_page_break()


def parrafo(documento: Document, texto: str) -> None:
    documento.add_paragraph(texto)


def lista_numerada(documento: Document, items: list[str]) -> None:
    for item in items:
        p = documento.add_paragraph(item, style="List Number")


def lista_viñetas(documento: Document, items: list[str]) -> None:
    for item in items:
        documento.add_paragraph(item, style="List Bullet")


def nota(documento: Document, titulo: str, texto: str) -> None:
    p = documento.add_paragraph()
    r1 = p.add_run(f"{titulo}: ")
    r1.bold = True
    r1.font.color.rgb = COLOR_ACENTO
    p.add_run(texto)


def tabla_menu(documento: Document, filas: list[tuple[str, str]]) -> None:
    tabla = documento.add_table(rows=1, cols=2)
    tabla.style = "Light Grid Accent 1"
    hdr = tabla.rows[0].cells
    hdr[0].text = "Menú"
    hdr[1].text = "Secciones disponibles"
    for menu, secciones in filas:
        row = tabla.add_row().cells
        row[0].text = menu
        row[1].text = secciones
    documento.add_paragraph()


def construir_manual() -> Document:
    doc = Document()
    configurar_estilos(doc)

    seccion = doc.sections[0]
    seccion.top_margin = Cm(2.5)
    seccion.bottom_margin = Cm(2.5)
    seccion.left_margin = Cm(2.8)
    seccion.right_margin = Cm(2.5)

    pie = seccion.footer.paragraphs[0]
    pie.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pie.add_run("ERP Tienda — Manual de usuario  |  Página ")
    agregar_campo_numero_pagina(pie)

    portada(doc)

    doc.add_heading("Índice", level=1)
    indice = [
        "1. Introducción",
        "2. Acceso al sistema",
        "3. Navegación general",
        "4. Configuración inicial del negocio",
        "5. Clientes",
        "6. Proveedores",
        "7. Productos y categorías",
        "8. Stock",
        "9. Ventas",
        "10. Compras",
        "11. Cuentas corrientes",
        "12. Reportes",
        "13. Usuarios y permisos",
        "14. Consejos y buenas prácticas",
        "15. Soporte",
    ]
    for linea in indice:
        doc.add_paragraph(linea)
    doc.add_page_break()

    doc.add_heading("1. Introducción", level=1)
    parrafo(
        doc,
        "ERP Tienda es un sistema de gestión pensado para comercios minoristas. Permite "
        "administrar clientes y proveedores, cargar productos, controlar el inventario, "
        "registrar ventas y compras, consultar cuentas corrientes y obtener reportes "
        "operativos desde una única plataforma web.",
    )
    parrafo(
        doc,
        "Este manual está dirigido al usuario final (encargado de mostrador, administración "
        "o dueño del negocio) y describe los pasos habituales de trabajo. No requiere "
        "conocimientos técnicos.",
    )

    doc.add_heading("2. Acceso al sistema", level=1)
    doc.add_heading("2.1 Inicio de sesión", level=2)
    lista_numerada(
        doc,
        [
            "Abra el navegador web (Chrome, Firefox o Edge recomendados).",
            "Ingrese la dirección web que le proporcionó el administrador del sistema.",
            "En la pantalla de inicio de sesión, complete su usuario y contraseña.",
            "Pulse Ingresar.",
        ],
    )
    nota(
        doc,
        "Importante",
        "La primera vez, el administrador del sistema le entregará sus credenciales. "
        "Cámbielas desde Usuarios → Alta usuarios en cuanto ingrese por primera vez.",
    )

    doc.add_heading("2.2 Cierre de sesión", level=2)
    parrafo(
        doc,
        "Utilice la opción de cerrar sesión en la barra superior cuando termine de trabajar, "
        "especialmente en equipos compartidos.",
    )

    doc.add_heading("3. Navegación general", level=1)
    parrafo(
        doc,
        "Una vez autenticado, verá el panel principal con un menú lateral (en computadoras) "
        "o una barra inferior (en celulares). Cada ítem del menú agrupa las tareas del día a día.",
    )
    tabla_menu(
        doc,
        [
            ("Inicio", "Resumen del día: ventas, alertas de stock y cuentas corrientes."),
            ("Clientes", "Alta clientes · Cuentas corrientes"),
            ("Ventas", "Centro de ventas · Historial de ventas"),
            ("Compras", "Proveedores · Registro de compras"),
            ("Productos", "Catálogo · Categorías"),
            ("Stock", "Stock actual · Auditorías de stock"),
            ("Reportes", "Informes de ventas, stock, compras y más."),
            ("Usuarios", "Alta usuarios · Permisos (solo perfiles autorizados)"),
            ("Configuración", "Datos del negocio · Parámetros del sistema"),
        ],
    )
    nota(
        doc,
        "Permisos",
        "Es posible que no vea todas las secciones: el administrador define qué menús "
        "puede usar cada usuario.",
    )

    doc.add_heading("4. Configuración inicial del negocio", level=1)
    parrafo(
        doc,
        "Antes de operar ventas, complete los datos del negocio en Configuración → Negocio:",
    )
    lista_viñetas(
        doc,
        [
            "Nombre comercial, razón social y datos de contacto.",
            "Logo (opcional) para comprobantes y reportes.",
            "Redes sociales visibles en documentos impresos.",
        ],
    )
    parrafo(
        doc,
        "En Configuración → Sistema puede ajustar parámetros como umbrales de stock bajo "
        "y otras preferencias operativas, según los permisos de su usuario.",
    )

    doc.add_heading("5. Clientes", level=1)
    doc.add_heading("5.1 Alta de un cliente", level=2)
    lista_numerada(
        doc,
        [
            "Menú Clientes → Alta clientes.",
            "Pulse el botón Nuevo cliente (parte superior).",
            "Complete los datos de identificación: nombre, documento y condición frente al IVA.",
            "Complete contacto: teléfonos, correo y dirección.",
            "Si el cliente opera con cuenta corriente, active la opción correspondiente e indique el límite de crédito.",
            "Pulse Crear cliente.",
        ],
    )
    doc.add_heading("5.2 Consultar o modificar un cliente", level=2)
    lista_numerada(
        doc,
        [
            "En la grilla de clientes, use el buscador para filtrar por nombre o documento.",
            "Pulse Detalle en la tarjeta del cliente.",
            "La ficha se abre en solo lectura. Pulse Editar para modificar datos.",
            "El interruptor Habilitado (encabezado) permite habilitar o inhabilitar ventas al cliente.",
            "Pulse Guardar cambios para confirmar.",
        ],
    )
    nota(
        doc,
        "Tip",
        "Un cliente inhabilitado no podrá seleccionarse en el centro de ventas.",
    )

    doc.add_heading("6. Proveedores", level=1)
    doc.add_heading("6.1 Alta de un proveedor", level=2)
    lista_numerada(
        doc,
        [
            "Menú Compras → Proveedores.",
            "Pulse Nuevo proveedor.",
            "Complete nombre, documento, datos de contacto y dirección.",
            "Si compra a crédito, active Compras a crédito e indique el tope permitido.",
            "Pulse Crear proveedor.",
        ],
    )
    doc.add_heading("6.2 Consultar o modificar", level=2)
    parrafo(
        doc,
        "El flujo es igual al de clientes: buscar en la grilla, abrir Detalle, Editar si "
        "corresponde y Guardar cambios. Los chips en cada tarjeta muestran si el proveedor "
        "está habilitado para compras y si tiene crédito.",
    )

    doc.add_heading("7. Productos y categorías", level=1)
    doc.add_heading("7.1 Categorías", level=2)
    parrafo(doc, "Menú Productos → Categorías. Cree las familias de productos (ej.: Remeras, Calzado) antes de cargar el catálogo.")
    doc.add_heading("7.2 Catálogo de productos", level=2)
    lista_numerada(
        doc,
        [
            "Menú Productos → Catálogo.",
            "Pulse Nuevo producto e ingrese nombre, marca, categoría, precio de venta y costo.",
            "Agregue variantes (talle, color, código de barras) según corresponda.",
            "Cada variante es la unidad que se vende y controla en stock.",
        ],
    )
    nota(
        doc,
        "Código de barras",
        "Asignar un código de barras a cada variante permite escanear productos en el centro de ventas.",
    )

    doc.add_heading("8. Stock", level=1)
    doc.add_heading("8.1 Consultar stock actual", level=2)
    parrafo(
        doc,
        "Menú Stock → Stock actual. Visualice cantidades por variante, filtre por categoría "
        "o producto y detecte faltantes o unidades bajas.",
    )
    doc.add_heading("8.2 Ajustes e ingresos manuales", level=2)
    parrafo(
        doc,
        "Desde Stock actual puede registrar ajustes por conteo físico o entradas manuales "
        "(según permisos). También puede importar un archivo de conteo si su operación lo utiliza.",
    )
    doc.add_heading("8.3 Auditorías", level=2)
    parrafo(
        doc,
        "Menú Stock → Auditorías de stock. Revise el historial de conteos y ajustes "
        "realizados, con fecha y usuario responsable.",
    )
    nota(
        doc,
        "Automático",
        "Las ventas descuentan stock y las compras lo incrementan automáticamente.",
    )

    doc.add_heading("9. Ventas", level=1)
    doc.add_heading("9.1 Centro de ventas", level=2)
    parrafo(doc, "Menú Ventas → Centro de ventas. Esta es la pantalla principal para cobrar en mostrador.")
    doc.add_heading("9.2 Agregar productos al ticket", level=3)
    lista_viñetas(
        doc,
        [
            "Modo Lector: escanee el código de barras o escriba el código y presione Enter.",
            "Modo Nombre: busque por nombre, marca o categoría y seleccione la variante.",
        ],
    )
    doc.add_heading("9.3 Cliente y forma de pago", level=3)
    lista_numerada(
        doc,
        [
            "Seleccione un cliente registrado desde el buscador, o elija Consumidor final.",
            "Para consumidor final puede opcionalmente cargar documento y nombre.",
            "Elija la forma de pago: efectivo, tarjeta, transferencia o cuenta corriente (si el cliente la tiene habilitada).",
            "Revise totales en el panel inferior.",
        ],
    )
    doc.add_heading("9.4 Confirmar la venta", level=3)
    lista_numerada(
        doc,
        [
            "Pulse Confirmar venta.",
            "Verá un mensaje de éxito con el número de comprobante.",
            "Puede imprimir o exportar el resumen de la venta en PDF desde la misma pantalla.",
        ],
    )
    doc.add_heading("9.5 Historial de ventas", level=2)
    parrafo(
        doc,
        "Menú Ventas → Historial de ventas. Consulte ventas anteriores, filtre por fecha "
        "y acceda al detalle de cada operación.",
    )

    doc.add_heading("10. Compras", level=1)
    doc.add_heading("10.1 Registrar una compra", level=2)
    lista_numerada(
        doc,
        [
            "Menú Compras → Registro de compras.",
            "Pulse Nueva compra.",
            "Seleccione proveedor y condición (contado o crédito, si aplica).",
            "Agregue líneas: variante, cantidad y costo unitario.",
            "Confirme el registro. El stock se actualizará automáticamente.",
        ],
    )

    doc.add_heading("11. Cuentas corrientes", level=1)
    parrafo(
        doc,
        "Menú Clientes → Cuentas corrientes. Desde aquí puede:",
    )
    lista_viñetas(
        doc,
        [
            "Ver saldos pendientes de cada cliente con cuenta corriente habilitada.",
            "Consultar movimientos (ventas a crédito y pagos).",
            "Registrar un pago: indique monto, forma de pago y observaciones.",
            "Imprimir estado de cuenta o recibo de pago.",
        ],
    )

    doc.add_heading("12. Reportes", level=1)
    parrafo(
        doc,
        "Menú Reportes. Seleccione el informe deseado, aplique filtros (fechas, categorías, "
        "proveedores, etc.) y genere la vista en pantalla. La mayoría permite exportar a PDF "
        "o Excel según el tipo de reporte.",
    )
    lista_viñetas(
        doc,
        [
            "Ventas por período y facturación.",
            "Stock valorizado, stock crítico y movimientos.",
            "Cuentas corrientes y compras por proveedor.",
            "Productos más vendidos y ventas por categoría.",
        ],
    )

    doc.add_heading("13. Usuarios y permisos", level=1)
    parrafo(
        doc,
        "Sección destinada al responsable o administrador del negocio (no a todos los operadores).",
    )
    doc.add_heading("13.1 Alta de usuarios", level=2)
    lista_numerada(
        doc,
        [
            "Menú Usuarios → Alta usuarios → Nuevo usuario.",
            "Complete nombre, apellido, usuario de login y contraseña inicial.",
            "Asigne rol: Dueño o Empleado (el rol Administrador es exclusivo del proveedor del sistema).",
            "Guarde la ficha.",
        ],
    )
    doc.add_heading("13.2 Permisos", level=2)
    parrafo(
        doc,
        "Menú Usuarios → Permisos usuario. Defina qué menús ve cada empleado y qué acciones "
        "puede realizar (ajustar stock, registrar compras, blanquear contraseñas, etc.).",
    )

    doc.add_heading("14. Consejos y buenas prácticas", level=1)
    lista_viñetas(
        doc,
        [
            "Mantenga el catálogo actualizado antes de vender: precios, variantes y códigos de barras.",
            "Revise el panel Inicio al abrir el turno para detectar stock bajo y deudas de clientes.",
            "Inhabilite clientes o proveedores en lugar de borrarlos si dejan de operar.",
            "Realice conteos periódicos de stock y registre auditorías en el sistema.",
            "Cierre sesión al terminar, especialmente en equipos compartidos.",
            "Ante un error de acceso denegado, consulte con quien administra permisos de usuario.",
        ],
    )

    doc.add_heading("15. Soporte", level=1)
    parrafo(
        doc,
        "Ante inconvenientes técnicos (pantalla en blanco, imposibilidad de ingresar, errores "
        "al guardar), anote el mensaje que aparece en pantalla y contacte al administrador "
        "del sistema o al soporte técnico que instaló la plataforma.",
    )
    parrafo(
        doc,
        "Indique siempre: usuario con el que operaba, menú en el que estaba, hora aproximada "
        "y pasos que realizó antes del error.",
    )

    return doc


def main() -> None:
    documento = construir_manual()
    documento.save(SALIDA)
    print(f"Manual generado: {SALIDA}")


if __name__ == "__main__":
    main()
